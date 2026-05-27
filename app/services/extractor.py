import io
import csv
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MAX_CHARS = 4000

_ocr_instance = None


def _get_ocr():
    global _ocr_instance
    if _ocr_instance is None:
        from paddleocr import PaddleOCR

        _ocr_instance = PaddleOCR(lang="ch")
    return _ocr_instance


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = Path(filename).suffix.lower()

    handlers = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".xls": _extract_xlsx,
        ".csv": _extract_csv,
        ".tsv": _extract_csv,
        ".txt": _extract_txt,
        ".md": _extract_txt,
        ".png": _extract_image,
        ".jpg": _extract_image,
        ".jpeg": _extract_image,
    }

    handler = handlers.get(ext)
    if handler is None:
        return f"[不支持的文件格式: {ext}]"

    try:
        text = handler(file_bytes)
        return text[:MAX_CHARS]
    except Exception as e:
        logger.warning(f"文本提取失败 ({filename}): {e}")
        return f"[文本提取失败: {e}]"


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    empty_pages = 0
    for page in reader.pages:
        t = page.extract_text()
        if t and t.strip():
            pages.append(t.strip())
        else:
            empty_pages += 1

    text_result = "\n\n".join(pages) if pages else ""

    total = len(reader.pages)
    if total > 0 and empty_pages / total > 0.5:
        logger.info(f"PDF 扫描件检测: {empty_pages}/{total} 页无文字层，启动 OCR")
        ocr_text = _ocr_pdf_pages(data)
        if ocr_text:
            return ocr_text

    return text_result if text_result else "[PDF 无可提取文本]"


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs) if paragraphs else "[DOCX 无可提取文本]"


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True)
    sheets = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | " + " | ".join(cells) + " |")
        if rows:
            sheets.append(f"## Sheet: {name}\n" + "\n".join(rows))
    return "\n\n".join(sheets) if sheets else "[XLSX 无可提取数据]"


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    delimiter = "\t" if "\t" in text[:200] else ","

    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = []
    for row in reader:
        if any(row):
            rows.append(" | " + " | ".join(str(c).strip() for c in row) + " |")
    return "\n".join(rows) if rows else "[CSV 无可提取数据]"


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _extract_image(data: bytes) -> str:
    result = _ocr_image(data)
    return result if result else "[图片无可识别文字]"


def _ocr_image(img_bytes: bytes) -> str:
    try:
        ocr = _get_ocr()
        import numpy as np
        from PIL import Image

        image = Image.open(io.BytesIO(img_bytes))
        img_array = np.array(image)
        result = ocr.ocr(img_array)

        if not result or not result[0]:
            return ""

        lines = []
        for line_info in result[0]:
            text = line_info[1][0]
            if text and text.strip():
                lines.append(text.strip())
        return "\n".join(lines)
    except Exception as e:
        logger.warning(f"OCR 识别失败: {e}")
        return ""


def _ocr_pdf_pages(data: bytes) -> str:
    try:
        import pymupdf

        doc = pymupdf.open(stream=data, filetype="pdf")
        all_text = []
        for i in range(len(doc)):
            page = doc[i]
            pix = page.get_pixmap(dpi=200)
            page_text = _ocr_image(pix.tobytes("png"))
            if page_text:
                all_text.append(f"[Page {i + 1}]\n{page_text}")
        doc.close()
        return "\n\n".join(all_text) if all_text else ""
    except Exception as e:
        logger.warning(f"PDF OCR 降级失败: {e}")
        return ""